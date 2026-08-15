import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches

st.set_page_config(
    page_title="Bitácora Digital - Diagnóstico, Gemelo Digital y Simulación",
    page_icon="📝",
    layout="wide"
)

# ------------------------------------------------------------------------------
# MÓDULO DE RECUPERACIÓN Y GUARDADO VÍA JSON
# ------------------------------------------------------------------------------
st.sidebar.title("💾 Gestión de Avance (JSON)")
st.sidebar.info(
    "Sube tu archivo .json previo para restaurar tu trabajo o descarga el estado "
    "actual para continuar desde casa."
)

archivo_json_cargado = st.sidebar.file_uploader(
    "📂 Cargar Avance Guardado (.json)", type=["json"], key="uploader_json"
)

if archivo_json_cargado is not None:
    try:
        datos_recuperados = json.load(archivo_json_cargado)
        for k, v in datos_recuperados.items():
            st.session_state[k] = v
        st.sidebar.success("¡Avance cargado exitosamente! 🎉")
    except Exception as e:
        st.sidebar.error("Error al procesar el archivo JSON.")

default_keys = {
    "est1_nombre": "",
    "est1_email": "",
    "est2_nombre": "",
    "est2_email": "",
    "ans_2_1": "",
    "ans_2_2": "",
    "ans_2_3": "",
    "cajeros_val": 1,
    "link_bpmn": "",
    "txt_flexsim": "",
    "diagnostico_equipo": "",
    "recomendaciones_equipo": "",
    "tipo_g1": "Barras Simples",
    # Nuevas llaves para la valoración del aprendizaje
    "eval_matematica": 3,
    "eval_flexsim": 3,
    "eval_triangulacion": 3,
    "eval_sistemico": 3,
    "reflexion_aprensizaje": "",
    "dificultades_superadas": "",
    "aplicaciones_futuras": ""
}
for k, v in default_keys.items():
    if k not in st.session_state:
        st.session_state[k] = v

st.title("🏦 Bitácora Digital de Consultoría - Diagnóstico Operativo y Gemelo Digital")
st.caption("Herramienta de Recolección de Datos, Mapeo, Parametrización, Simulación en FlexSim, Resultados y Triangulación")
st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 1: REGISTRO DEL EQUIPO
# ------------------------------------------------------------------------------
st.header("📋 Bloque 1: Datos del Equipo Consultor")
col_e1, col_e2 = st.columns(2)

with col_e1:
    st.subheader("Estudiante 1")
    estudiante1 = st.text_input(
        "Nombre Estudiante 1",
        value=st.session_state["est1_nombre"],
        placeholder="Ingrese nombre completo...",
        key="est1_nombre"
    )
    email1 = st.text_input(
        "Correo Electrónico 1",
        value=st.session_state["est1_email"],
        placeholder="ejemplo@correo.com",
        key="est1_email"
    )

with col_e2:
    st.subheader("Estudiante 2")
    estudiante2 = st.text_input(
        "Nombre Estudiante 2",
        value=st.session_state["est2_nombre"],
        placeholder="Ingrese nombre completo...",
        key="est2_nombre"
    )
    email2 = st.text_input(
        "Correo Electrónico 2",
        value=st.session_state["est2_email"],
        placeholder="ejemplo@correo.com",
        key="est2_email"
    )

st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 2: MATRICES DE DATOS, GRÁFICAS Y ANÁLISIS DE CAMPO (AS-IS)
# ------------------------------------------------------------------------------
st.header("📊 Bloque 2: Digitación de Datos, Visualización y Diagnóstico Operativo (As-Is)")

# PASO 2.1
st.subheader("Paso 2.1: Mezcla de Trámites, Tiempos de Atención y Autorización")

df_tramites_init = pd.DataFrame({
    "ID": [1, 2, 3, 4],
    "Tipo_Tramite": ["Tipico", "Pesada", "Largos", ""],
    "Mezcla_Pct": [50.0, 30.0, 20.0, 0.0],
    "Tiempo_Atencion_Seg": [120.0, 400.0, 2000.0, 0.0],
    "Prob_Autorizacion": [0.0, 30.0, 35.0, 0.0]
})
df_tramites = st.data_editor(df_tramites_init, num_rows="dynamic", use_container_width=True, key="ed_tramites")

df_t_valid = df_tramites[df_tramites['Tipo_Tramite'] != ""].copy()

suma_mezcla = df_t_valid['Mezcla_Pct'].sum()
if not df_t_valid.empty and abs(suma_mezcla - 100.0) > 0.01:
    st.warning(f"⚠️ Atención: La suma de la mezcla de trámites actual es **{suma_mezcla:.1f}%**. Debe sumar el **100%** para un análisis correcto.")

buf_fig1_tiempos = None
buf_fig1_prob = None

if not df_t_valid.empty and df_t_valid['Tiempo_Atencion_Seg'].sum() > 0:
    col_g1, col_g2 = st.columns(2)
    
    with col_g1:
        tipo_grafica_1 = st.radio(
            "Seleccione el tipo de gráfica para Tiempos:",
            ["Barras Simples", "Diagrama de Pareto"],
            horizontal=True,
            key="tipo_g1"
        )
        
        df_bars = df_t_valid.sort_values(by='Tiempo_Atencion_Seg', ascending=False).reset_index(drop=True)
        etiquetas_x = [f"{row['Tipo_Tramite']}\n({row['Mezcla_Pct']:.0f}%)" for _, row in df_bars.iterrows()]
        
        fig_t, ax_t = plt.subplots(figsize=(5, 3.8))
        
        if tipo_grafica_1 == "Barras Simples":
            bars1 = ax_t.bar(etiquetas_x, df_bars['Tiempo_Atencion_Seg'], color='steelblue', alpha=0.8, width=0.4)
            for bar, (_, row) in zip(bars1, df_bars.iterrows()):
                yval = bar.get_height()
                pct = row['Mezcla_Pct']
                ax_t.text(
                    bar.get_x() + bar.get_width()/2,
                    yval + (df_bars['Tiempo_Atencion_Seg'].max() * 0.02),
                    f"{yval:.0f}s ({pct:.0f}%)",
                    ha='center', va='bottom', fontsize=8.5, fontweight='bold'
                )

            ax_t.set_ylabel('Tiempo Atención (seg)', fontweight='bold')
            ax_t.set_xlabel('Trámite (% Mezcla)', fontweight='bold')
            ax_t.set_ylim(0, df_bars['Tiempo_Atencion_Seg'].max() * 1.18)
            plt.xticks(rotation=0)
            plt.title("Tiempo de Atención y Participación (%) por Trámite", fontsize=9.5, fontweight='bold')
            
        else:
            df_bars['Acumulado_Pct'] = (df_bars['Tiempo_Atencion_Seg'].cumsum() / df_bars['Tiempo_Atencion_Seg'].sum()) * 100
            
            bars1 = ax_t.bar(etiquetas_x, df_bars['Tiempo_Atencion_Seg'], color='steelblue', alpha=0.8, width=0.4)
            ax_t.set_ylabel('Tiempo Atención (seg)', fontweight='bold', color='steelblue')
            ax_t.tick_params(axis='y', labelcolor='steelblue')
            
            for bar, (_, row) in zip(bars1, df_bars.iterrows()):
                yval = bar.get_height()
                pct = row['Mezcla_Pct']
                ax_t.text(
                    bar.get_x() + bar.get_width()/2,
                    yval + (df_bars['Tiempo_Atencion_Seg'].max() * 0.02),
                    f"{yval:.0f}s ({pct:.0f}%)",
                    ha='center', va='bottom', fontsize=8, fontweight='bold'
                )
            
            ax_t2 = ax_t.twinx()
            ax_t2.plot(etiquetas_x, df_bars['Acumulado_Pct'], color='crimson', marker='D', ms=5, linewidth=2)
            ax_t2.set_ylabel('% Acumulado Tiempos', fontweight='bold', color='crimson')
            ax_t2.tick_params(axis='y', labelcolor='crimson')
            ax_t2.set_ylim(0, 115)
            ax_t2.axhline(80, color='gray', linestyle='--', alpha=0.7)
            
            plt.xticks(rotation=0)
            plt.title("Pareto de Tiempos con Mezcla (%)", fontsize=9.5, fontweight='bold')

        fig_t.tight_layout()
        st.pyplot(fig_t)

        buf_fig1_tiempos = BytesIO()
        fig_t.savefig(buf_fig1_tiempos, format="png")
        buf_fig1_tiempos.seek(0)

    with col_g2:
        st.markdown("**Porcentaje de Autorizaciones:**")
        fig_pr, ax_pr = plt.subplots(figsize=(5, 3.8))
        bars2 = ax_pr.bar(df_t_valid['Tipo_Tramite'], df_t_valid['Prob_Autorizacion'], color='darkseagreen', width=0.4)
        
        for bar in bars2:
            yval = bar.get_height()
            ax_pr.text(bar.get_x() + bar.get_width()/2, yval + 1, f"{yval:.1f}%", ha='center', va='bottom', fontsize=9)

        ax_pr.set_ylabel('Probabilidad de Autorización (%)', fontweight='bold')
        ax_pr.set_xlabel('Tipo de Trámite', fontweight='bold')
        ax_pr.set_ylim(0, max(df_t_valid['Prob_Autorizacion'].max() + 15, 100))
        plt.xticks(rotation=15)
        plt.title("Porcentaje de Autorización", fontsize=10, fontweight='bold')
        
        fig_pr.tight_layout()
        st.pyplot(fig_pr)

        buf_fig1_prob = BytesIO()
        fig_pr.savefig(buf_fig1_prob, format="png")
        buf_fig1_prob.seek(0)

    st.markdown("#### 🔍 Análisis Diagnóstico (Paso 2.1)")
    analisis_2_1 = st.text_area(
        "Redacte su diagnóstico sobre los tiempos de atención y autorizaciones de los trámites:",
        value=st.session_state["ans_2_1"],
        height=100,
        key="ans_2_1"
    )

else:
    st.info("ℹ️ Llene la tabla superior para generar las gráficas de Tiempos y Autorizaciones.")

st.markdown("---")

# ------------------------------------------------------------------------------
# PASO 2.2: ANÁLISIS DE DEMANDA POR TIPO DE TRÁMITE Y COMPARATIVO DE PARETOS
# ------------------------------------------------------------------------------
st.subheader("Paso 2.2: Demanda y Análisis Pareto por Tipo de Trámite (Volumen vs Carga de Trabajo)")

df_arribos_init = pd.DataFrame({
    "Franja_Horaria": ["8:00 - 9:00", "9:00 - 10:00", "10:00 - 11:00", "11:00 - 12:00", "12:00 - 13:00", "13:00 - 14:00"],
    "Tipo_Perfil": ["Valle", "Pico", "Pico", "Valle", "Valle", "Valle"],
    "Clientes_Hora_Lambda": [15, 45, 50, 20, 15, 10]
})
df_arribos = st.data_editor(df_arribos_init, num_rows="dynamic", use_container_width=True, key="ed_arribos")

df_a_valid = df_arribos[df_arribos['Franja_Horaria'] != ""].copy()
total_clientes_dia = df_a_valid['Clientes_Hora_Lambda'].sum() if not df_a_valid.empty else 0

buf_fig_pareto_tramites = None
buf_fig_pareto_tiempo = None
buf_fig_tramite_hora = None

col_p1, col_p2 = st.columns([1.3, 1])

with col_p1:
    if not df_t_valid.empty and total_clientes_dia > 0:
        # 1. Datos para Pareto por Volumen de Clientes
        df_demanda_tramites = df_t_valid.copy()
        df_demanda_tramites['Clientes_Dia'] = (df_demanda_tramites['Mezcla_Pct'] / 100.0) * total_clientes_dia
        df_demanda_vol = df_demanda_tramites.sort_values(by='Clientes_Dia', ascending=False).reset_index(drop=True)
        df_demanda_vol['Acumulado_Pct'] = (df_demanda_vol['Clientes_Dia'].cumsum() / df_demanda_vol['Clientes_Dia'].sum()) * 100.0

        # 2. Datos para Pareto por Carga de Tiempo de Atención Total
        df_demanda_tramites['Tiempo_Total_Dia_Seg'] = df_demanda_tramites['Clientes_Dia'] * df_demanda_tramites['Tiempo_Atencion_Seg']
        df_demanda_tiempo = df_demanda_tramites.sort_values(by='Tiempo_Total_Dia_Seg', ascending=False).reset_index(drop=True)
        df_demanda_tiempo['Acumulado_Pct_Tiempo'] = (df_demanda_tiempo['Tiempo_Total_Dia_Seg'].cumsum() / df_demanda_tiempo['Tiempo_Total_Dia_Seg'].sum()) * 100.0

        tab_p_dia, tab_p_tiempo, tab_p_hora = st.tabs([
            "📊 Pareto Volumen Clientes (Cortos)", 
            "⏳ Pareto Carga de Tiempo (Largos)", 
            "🕒 Trámites por Hora"
        ])

        # TAB 1: PARETO VOLUMEN
        with tab_p_dia:
            fig_pt, ax_pt1 = plt.subplots(figsize=(5.5, 3.8))
            bars_pt = ax_pt1.bar(df_demanda_vol['Tipo_Tramite'], df_demanda_vol['Clientes_Dia'], color='teal', alpha=0.8, width=0.4)

            for bar in bars_pt:
                yval = bar.get_height()
                ax_pt1.text(bar.get_x() + bar.get_width()/2, yval + (df_demanda_vol['Clientes_Dia'].max() * 0.02), f"{yval:.0f} pers", ha='center', va='bottom', fontsize=8, fontweight='bold')

            ax_pt1.set_ylabel('Total Clientes / Día', fontweight='bold', color='teal')
            ax_pt1.tick_params(axis='y', labelcolor='teal')

            ax_pt2 = ax_pt1.twinx()
            ax_pt2.plot(df_demanda_vol['Tipo_Tramite'], df_demanda_vol['Acumulado_Pct'], color='crimson', marker='D', ms=5, linewidth=2)
            ax_pt2.set_ylabel('% Acumulado Personas', fontweight='bold', color='crimson')
            ax_pt2.tick_params(axis='y', labelcolor='crimson')
            ax_pt2.set_ylim(0, 115)
            ax_pt2.axhline(80, color='gray', linestyle='--', alpha=0.7)

            plt.title(f"Pareto por Volumen de Personas (Saturación de Sala)", fontsize=8.5, fontweight='bold')
            fig_pt.tight_layout()
            st.pyplot(fig_pt)

            buf_fig_pareto_tramites = BytesIO()
            fig_pt.savefig(buf_fig_pareto_tramites, format="png")
            buf_fig_pareto_tramites.seek(0)

        # TAB 2: PARETO TIEMPO
        with tab_p_tiempo:
            fig_time, ax_tm1 = plt.subplots(figsize=(5.5, 3.8))
            horas_trabajo = df_demanda_tiempo['Tiempo_Total_Dia_Seg'] / 3600.0
            bars_tm = ax_tm1.bar(df_demanda_tiempo['Tipo_Tramite'], horas_trabajo, color='darkorange', alpha=0.8, width=0.4)

            for bar in bars_tm:
                yval = bar.get_height()
                ax_tm1.text(bar.get_x() + bar.get_width()/2, yval + (horas_trabajo.max() * 0.02), f"{yval:.1f} hrs", ha='center', va='bottom', fontsize=8, fontweight='bold')

            ax_tm1.set_ylabel('Horas de Atención Requeridas / Día', fontweight='bold', color='darkorange')
            ax_tm1.tick_params(axis='y', labelcolor='darkorange')

            ax_tm2 = ax_tm1.twinx()
            ax_tm2.plot(df_demanda_tiempo['Tipo_Tramite'], df_demanda_tiempo['Acumulado_Pct_Tiempo'], color='crimson', marker='D', ms=5, linewidth=2)
            ax_tm2.set_ylabel('% Acumulado Tiempo Ocupación', fontweight='bold', color='crimson')
            ax_tm2.tick_params(axis='y', labelcolor='crimson')
            ax_tm2.set_ylim(0, 115)
            ax_tm2.axhline(80, color='gray', linestyle='--', alpha=0.7)

            plt.title("Pareto por Carga de Tiempo de Atención (Cuello de Botella)", fontsize=8.5, fontweight='bold')
            fig_time.tight_layout()
            st.pyplot(fig_time)

            buf_fig_pareto_tiempo = BytesIO()
            fig_time.savefig(buf_fig_pareto_tiempo, format="png")
            buf_fig_pareto_tiempo.seek(0)

        # TAB 3: TRÁMITES POR HORA
        with tab_p_hora:
            fig_th, ax_th = plt.subplots(figsize=(5.5, 3.8))
            bottom_stack = np.zeros(len(df_a_valid))
            
            for _, row_t in df_t_valid.iterrows():
                tramite_nom = row_t['Tipo_Tramite']
                pct_t = row_t['Mezcla_Pct'] / 100.0
                valores_h = df_a_valid['Clientes_Hora_Lambda'] * pct_t
                ax_th.bar(df_a_valid['Franja_Horaria'], valores_h, bottom=bottom_stack, label=tramite_nom, width=0.45)
                bottom_stack += valores_h

            ax_th.set_ylabel('Clientes / Hora', fontweight='bold')
            ax_th.set_xlabel('Franja Horaria', fontweight='bold')
            plt.xticks(rotation=25)
            ax_th.legend(title="Trámites", fontsize=8)
            plt.title("Estimación de Arribos por Trámite en cada Franja Horaria", fontsize=9, fontweight='bold')
            fig_th.tight_layout()
            st.pyplot(fig_th)

            buf_fig_tramite_hora = BytesIO()
            fig_th.savefig(buf_fig_tramite_hora, format="png")
            buf_fig_tramite_hora.seek(0)

    else:
        st.info("ℹ️ Complete los datos de Trámites (Paso 2.1) y Arribos para generar las gráficas Pareto de trámites.")

with col_p2:
    st.markdown("#### 🔍 Análisis Diagnóstico de Demanda (Paso 2.2)")
    analisis_2_2 = st.text_area(
        "Redacte su análisis contrastando la saturación por volumen (trámites cortos) vs. el cuello de botella por ocupación (trámites largos):",
        value=st.session_state["ans_2_2"],
        height=220,
        key="ans_2_2"
    )

st.markdown("---")

# ------------------------------------------------------------------------------
# PASO 2.3: CAPACIDAD DEL SISTEMA VS ARRIBOS
# ------------------------------------------------------------------------------
st.subheader("Paso 2.3: Curva de Capacidad del Sistema vs. Arribos")

cajeros = st.number_input("Número de Cajeros Activos en Ventanilla (c):", min_value=1, max_value=20, value=int(st.session_state["cajeros_val"]), key="cajeros_val")

col_c1, col_c2 = st.columns([1.3, 1])

buf_fig_capacidad = None

if not df_t_valid.empty:
    ts_ponderado_calc = ((df_t_valid['Mezcla_Pct'] / 100) * df_t_valid['Tiempo_Atencion_Seg']).sum()
else:
    ts_ponderado_calc = 0

with col_c1:
    if not df_a_valid.empty and df_a_valid['Clientes_Hora_Lambda'].sum() > 0:
        fig_cap, ax_c = plt.subplots(figsize=(5.5, 3.8))
        ax_c.plot(df_a_valid['Franja_Horaria'], df_a_valid['Clientes_Hora_Lambda'], marker='o', color='navy', linewidth=2.5, label='Llegadas (λ)')
        ax_c.fill_between(df_a_valid['Franja_Horaria'], df_a_valid['Clientes_Hora_Lambda'], color='skyblue', alpha=0.3)

        if ts_ponderado_calc > 0:
            capacidad_sis_calc = (3600 / ts_ponderado_calc) * cajeros
            ax_c.axhline(y=capacidad_sis_calc, color='red', linestyle='--', linewidth=2, label=f'Capacidad ({cajeros} cajeros)')

        ax_c.set_title("Arribos vs. Umbral de Capacidad Operativa", fontsize=10, fontweight='bold')
        ax_c.set_xlabel("Franja Horaria", fontweight='bold')
        ax_c.set_ylabel("Clientes / Hora", fontweight='bold')
        plt.xticks(rotation=25)
        ax_c.legend()
        ax_c.grid(True, linestyle=':', alpha=0.6)
        fig_cap.tight_layout()
        st.pyplot(fig_cap)
        
        buf_fig_capacidad = BytesIO()
        fig_cap.savefig(buf_fig_capacidad, format="png")
        buf_fig_capacidad.seek(0)
    else:
        st.info("ℹ️ Llene las franjas de arribos para visualizar la curva de capacidad.")

with col_c2:
    st.markdown("#### 🔍 Análisis Diagnóstico de Capacidad (Paso 2.3)")
    analisis_2_3 = st.text_area(
        "Redacte su evaluación sobre si el número de cajeros es suficiente para soportar la tasa de llegadas:",
        value=st.session_state.get("ans_2_3", ""),
        height=200,
        key="ans_2_3"
    )

st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 3: CONSTRUCCIÓN DEL GEMELO DIGITAL (FLEXSIM)
# ------------------------------------------------------------------------------
st.header("🖥️ Bloque 3: Construcción del Gemelo Digital (FlexSim)")

st.subheader("3.1 Definición del Proceso")
col_b1, col_b2 = st.columns([1, 1])
with col_b1:
    link_bpmn = st.text_input("🔗 Link del Diagrama BPMN:", value=st.session_state["link_bpmn"], key="link_bpmn")
    imagen_bpmn = st.file_uploader("🖼️ Cargar Pantallazo BPMN:", type=["png", "jpg", "jpeg"], key="up_bpmn")
with col_b2:
    if imagen_bpmn: st.image(imagen_bpmn, use_container_width=True)

df_bpmn_init = pd.DataFrame({"Etapa_ID": [1,2,3], "Nombre_Etapa": ["","",""], "Actor_Responsable": ["","",""], "Tipo_Actividad": ["","",""], "Descripción_Operativa": ["","",""]})
df_bpmn = st.data_editor(df_bpmn_init, num_rows="dynamic", use_container_width=True, key="ed_bpmn")
st.markdown("---")

st.subheader("3.2 Parámetros del Gemelo Digital")
df_params_init = pd.DataFrame({"Etapa_ID": [1,2,3], "Nombre_Etapa": ["","",""], "Objeto_FlexSim": ["","",""], "Distribución_o_Regla": ["","",""], "Parámetros_Numéricos": ["","",""], "Recursos_Asignados": ["","",""]})
df_params = st.data_editor(df_params_init, num_rows="dynamic", use_container_width=True, key="ed_params")
st.markdown("---")

st.subheader("3.3 Modelo Construido en FlexSim")
col_f1, col_f2 = st.columns([1, 1])
with col_f1:
    imagen_flexsim = st.file_uploader("🖼️ Cargar Pantallazo del Layout FlexSim:", type=["png", "jpg", "jpeg"], key="up_flexsim")
    notas_flexsim = st.text_area("Observaciones sobre el modelo en FlexSim:", value=st.session_state["txt_flexsim"], height=150, key="txt_flexsim")
with col_f2:
    if imagen_flexsim: st.image(imagen_flexsim, use_container_width=True)
st.markdown("---")

st.subheader("3.4 Indicadores de Desempeño a Medir")
df_kpi_init = pd.DataFrame({"KPI_ID": [1,2,3], "Nombre_Indicador": ["","",""], "Unidad_Medida": ["","",""], "Objetivo": ["Minimizar","Minimizar","Minimizar"], "Meta_Verde": [0.0,0.0,0.0], "Critico_Rojo": [0.0,0.0,0.0]})
df_kpis = st.data_editor(df_kpi_init, num_rows="dynamic", use_container_width=True, column_config={"Objetivo": st.column_config.SelectboxColumn(options=["Minimizar", "Maximizar"])}, key="ed_kpis")
st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 4: SIMULACIÓN Y EVIDENCIAS
# ------------------------------------------------------------------------------
st.header("🎬 Bloque 4: Simulación y Evidencias de Ejecución (FlexSim)")
col_s1, col_s2 = st.columns(2)
with col_s1:
    imagen_sim_corrida = st.file_uploader("🖼️ Pantallazo de Simulación Ejecutándose:", type=["png", "jpg", "jpeg"], key="up_sim_corrida")
    if imagen_sim_corrida: st.image(imagen_sim_corrida, use_container_width=True)
with col_s2:
    imagen_dashboard_kpi = st.file_uploader("🖼️ Pantallazo del Dashboard de KPIs:", type=["png", "jpg", "jpeg"], key="up_dash_kpi")
    if imagen_dashboard_kpi: st.image(imagen_dashboard_kpi, use_container_width=True)
st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 5: TRIANGULACIÓN, DIAGNÓSTICO Y RECOMENDACIONES
# ------------------------------------------------------------------------------
st.header("⚖️ Bloque 5: Triangulación, Diagnóstico y Recomendaciones")

df_kpis_valid = df_kpis[df_kpis['Nombre_Indicador'] != ""].copy()
if not df_kpis_valid.empty:
    df_triang_init = pd.DataFrame({"KPI_ID": df_kpis_valid["KPI_ID"].tolist(), "Indicador_Clave": df_kpis_valid["Nombre_Indicador"].tolist(), "Modelo_Teorico_Calculado": [0.0]*len(df_kpis_valid), "Simulacion_FlexSim_Obtenido": [0.0]*len(df_kpis_valid)})
else:
    df_triang_init = pd.DataFrame({"KPI_ID": [1,2], "Indicador_Clave": ["",""], "Modelo_Teorico_Calculado": [0.0, 0.0], "Simulacion_FlexSim_Obtenido": [0.0, 0.0]})

df_triang_input = st.data_editor(df_triang_init, use_container_width=True, disabled=["KPI_ID", "Indicador_Clave"], key="ed_triangulacion_input")

resultados_triang = []
for i, row in df_triang_input.iterrows():
    calc, sim = row["Modelo_Teorico_Calculado"], row["Simulacion_FlexSim_Obtenido"]
    desv = sim - calc
    estado = "⚪ N/A"
    if not df_kpis_valid.empty and row["Indicador_Clave"] != "":
        kpi_info = df_kpis_valid[df_kpis_valid["KPI_ID"] == row["KPI_ID"]]
        if not kpi_info.empty:
            kpi_info = kpi_info.iloc[0]
            if kpi_info["Objetivo"] == "Minimizar":
                estado = "🟢 Verde" if sim <= kpi_info["Meta_Verde"] else ("🔴 Rojo" if sim >= kpi_info["Critico_Rojo"] else "🟡 Amarillo")
            else:
                estado = "🟢 Verde" if sim >= kpi_info["Meta_Verde"] else ("🔴 Rojo" if sim <= kpi_info["Critico_Rojo"] else "🟡 Amarillo")
    resultados_triang.append({"KPI_ID": row["KPI_ID"], "Indicador_Clave": row["Indicador_Clave"], "Modelo_Teorico_Calculado": calc, "Simulacion_FlexSim_Obtenido": sim, "Desviacion_o_Diferencia": desv, "Estado_Semaforo": estado})

df_triangulacion = pd.DataFrame(resultados_triang)
st.dataframe(df_triangulacion, use_container_width=True)

st.markdown("#### Diagnóstico y Recomendaciones")
diagnostico_equipo = st.text_area("Párrafo de Diagnóstico Consolidado:", value=st.session_state["diagnostico_equipo"], height=150, key="diagnostico_equipo")
recomendaciones_equipo = st.text_area("Propuestas y Recomendaciones de Intervención:", value=st.session_state["recomendaciones_equipo"], height=150, key="recomendaciones_equipo")
st.markdown("---")

# ------------------------------------------------------------------------------
# BLOQUE 6: VALORACIÓN DEL APRENDIZAJE Y CIERRE DE COMPETENCIAS (NUEVO)
# ------------------------------------------------------------------------------
st.header("🎓 Bloque 6: Valoración del Aprendizaje y Cierre de Competencias")
st.markdown(
    "Esta sección permite al equipo reflexionar sobre el proceso de consultoría, "
    "autoevaluar el desarrollo de sus competencias técnicas y consolidar las lecciones aprendidas."
)

st.subheader("6.1 Autoevaluación de Competencias Adquiridas")
st.caption("Califique en una escala de 1 (Inicial) a 5 (Dominio Alto) el nivel alcanzado en cada competencia:")

col_comp1, col_comp2 = st.columns(2)

with col_comp1:
    eval_matematica = st.slider(
        "1. Modelamiento Matemático y Analítico (Teoría de Colas, Variabilidad y Capacidad):",
        min_value=1, max_value=5, value=int(st.session_state["eval_matematica"]), key="eval_matematica"
    )
    eval_flexsim = st.slider(
        "2. Construcción de Gemelo Digital (Layout, Reglas de Ruteo y Parametrización en FlexSim):",
        min_value=1, max_value=5, value=int(st.session_state["eval_flexsim"]), key="eval_flexsim"
    )

with col_comp2:
    eval_triangulacion = st.slider(
        "3. Triangulación e Interpretación de Datos (Modelo Teórico vs. Simulación Dinámica):",
        min_value=1, max_value=5, value=int(st.session_state["eval_triangulacion"]), key="eval_triangulacion"
    )
    eval_sistemico = st.slider(
        "4. Pensamiento Sistémico y Propuesta de Soluciones (Diseño de Escenarios To-Be):",
        min_value=1, max_value=5, value=int(st.session_state["eval_sistemico"]), key="eval_sistemico"
    )

st.subheader("6.2 Preguntas de Cierre y Reflexión Metacognitiva")

reflexion_aprensizaje = st.text_area(
    "💬 1. Principales Aprendizajes: ¿Cuál fue el hallazgo más sorprendente al comparar los cálculos teóricos promedio frente al Gemelo Digital en FlexSim?",
    value=st.session_state["reflexion_aprensizaje"],
    height=120,
    key="reflexion_aprensizaje"
)

dificultades_superadas = st.text_area(
    "🛠️ 2. Retos y Dificultades: ¿Qué obstáculos técnicos o conceptuales enfrentó el equipo durante el modelamiento y cómo los superaron?",
    value=st.session_state["dificultades_superadas"],
    height=120,
    key="dificultades_superadas"
)

aplicaciones_futuras = st.text_area(
    "🚀 3. Transferencia y Aplicación: ¿Cómo aplicaría esta metodología de Gemelo Digital en otros procesos industriales o de servicios?",
    value=st.session_state["aplicaciones_futuras"],
    height=120,
    key="aplicaciones_futuras"
)

st.markdown("---")

# ------------------------------------------------------------------------------
# BOTÓN DE GUARDADO EN JSON
# ------------------------------------------------------------------------------
datos_a_guardar = {
    "est1_nombre": st.session_state.get("est1_nombre", ""),
    "est1_email": st.session_state.get("est1_email", ""),
    "est2_nombre": st.session_state.get("est2_nombre", ""),
    "est2_email": st.session_state.get("est2_email", ""),
    "ans_2_1": st.session_state.get("ans_2_1", ""),
    "ans_2_2": st.session_state.get("ans_2_2", ""),
    "ans_2_3": st.session_state.get("ans_2_3", ""),
    "cajeros_val": st.session_state.get("cajeros_val", 1),
    "link_bpmn": st.session_state.get("link_bpmn", ""),
    "txt_flexsim": st.session_state.get("txt_flexsim", ""),
    "diagnostico_equipo": st.session_state.get("diagnostico_equipo", ""),
    "recomendaciones_equipo": st.session_state.get("recomendaciones_equipo", ""),
    "tipo_g1": st.session_state.get("tipo_g1", "Barras Simples"),
    # Guardar respuestas del Bloque 6
    "eval_matematica": st.session_state.get("eval_matematica", 3),
    "eval_flexsim": st.session_state.get("eval_flexsim", 3),
    "eval_triangulacion": st.session_state.get("eval_triangulacion", 3),
    "eval_sistemico": st.session_state.get("eval_sistemico", 3),
    "reflexion_aprensizaje": st.session_state.get("reflexion_aprensizaje", ""),
    "dificultades_superadas": st.session_state.get("dificultades_superadas", ""),
    "aplicaciones_futuras": st.session_state.get("aplicaciones_futuras", "")
}
json_data = json.dumps(datos_a_guardar, indent=4, ensure_ascii=False)

st.sidebar.download_button(
    label="💾 Descargar Estado Actual (.json)",
    data=json_data,
    file_name="avance_bitacora2.json",
    mime="application/json"
)

# ------------------------------------------------------------------------------
# BLOQUE 7: GENERACIÓN DE DOCUMENTO WORD CONSOLIDADO
# ------------------------------------------------------------------------------
st.header("📄 Exportación de Informe Final")
st.info("Haga clic en el botón inferior para generar y descargar un documento de Word (.docx) que consolida todas las tablas, gráficas, diagnósticos y la autoevaluación del aprendizaje.")

def add_df_to_doc(df, doc):
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'Table Grid'
    for j in range(df.shape[-1]):
        t.cell(0,j).text = str(df.columns[j])
        t.cell(0,j).paragraphs[0].runs[0].font.bold = True
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

def generar_word():
    doc = Document()
    doc.add_heading('Informe Técnico de Consultoría - Diagnóstico Operativo y Gemelo Digital', 0)
    
    p_equipo = doc.add_paragraph()
    p_equipo.add_run('Equipo Consultor:\n').bold = True
    
    nom1 = estudiante1 if estudiante1 else "No registrado"
    em1 = f" ({email1})" if email1 else ""
    p_equipo.add_run(f'• Estudiante 1: {nom1}{em1}\n')
    
    nom2 = estudiante2 if estudiante2 else "No registrado"
    em2 = f" ({email2})" if email2 else ""
    p_equipo.add_run(f'• Estudiante 2: {nom2}{em2}')
    
    # SECCIÓN 1
    doc.add_heading('1. Diagnóstico Operativo del Sistema Actual (As-Is)', level=1)
    doc.add_paragraph("Para comprender la situación actual del sistema, se analizó inicialmente el catálogo de servicios. La Tabla 1 consolida los tipos de trámites identificados, su participación porcentual ideal de mezcla (cuya suma equivale al 100% de la demanda), el tiempo estimado y el porcentaje de autorización.")
    
    doc.add_heading('Tabla 1. Mezcla de Trámites, Tiempos Estándar y Autorizaciones', level=2)
    add_df_to_doc(df_tramites[df_tramites['Tipo_Tramite'] != ""], doc)
    
    if buf_fig1_tiempos:
        doc.add_paragraph("\nLa Figura 1 presenta visualmente los tiempos de atención por trámite integrando la proporción porcentual (%) que cada uno representa dentro de la mezcla total.")
        doc.add_picture(buf_fig1_tiempos, width=Inches(5.0))
        doc.add_paragraph(f"Figura 1. Tiempos de Atención y Porcentaje de Mezcla ({st.session_state.get('tipo_g1', 'Barras Simples')}).", style='Caption')

    if buf_fig1_prob:
        doc.add_paragraph("\nLa Figura 2 detalla la tasa o probabilidad de autorización correspondiente a cada tipo de trámite.")
        doc.add_picture(buf_fig1_prob, width=Inches(5.0))
        doc.add_paragraph("Figura 2. Porcentaje de Autorización por Tipo de Trámite.", style='Caption')
    
    p_ans1 = doc.add_paragraph()
    p_ans1.add_run("Análisis Diagnóstico de Trámites: ").bold = True
    p_ans1.add_run(analisis_2_1 if analisis_2_1 else "No se registraron observaciones.")
    
    doc.add_paragraph("\nPosteriormente, evaluando el flujo real de usuarios, la Tabla 2 detalla el perfil de arribos segmentado por franjas horarias.")
    doc.add_heading('Tabla 2. Perfil Horario de Arribos y Demanda', level=2)
    add_df_to_doc(df_arribos[df_arribos['Franja_Horaria'] != ""], doc)

    if buf_fig_pareto_tramites:
        doc.add_paragraph("\nLa Figura 3 refleja el Diagrama de Pareto por volumen de personas (identificando cómo los trámites cortos generan la saturación física de la sala).")
        doc.add_picture(buf_fig_pareto_tramites, width=Inches(5.0))
        doc.add_paragraph("Figura 3. Diagrama de Pareto de Clientes por Volumen de Trámite.", style='Caption')

    if buf_fig_pareto_tiempo:
        doc.add_paragraph("\nLa Figura 4 detalla el Diagrama de Pareto por Carga de Tiempo de Atención (evidenciando que los trámites largos son el principal cuello de botella del personal de ventanilla).")
        doc.add_picture(buf_fig_pareto_tiempo, width=Inches(5.0))
        doc.add_paragraph("Figura 4. Diagrama de Pareto de Carga de Trabajo (Horas de Atención).", style='Caption')

    if buf_fig_tramite_hora:
        doc.add_paragraph("\nLa Figura 5 ilustra la distribución acumulada por hora de los tipos de trámites solicitados.")
        doc.add_picture(buf_fig_tramite_hora, width=Inches(5.0))
        doc.add_paragraph("Figura 5. Composición de Trámites por Franja Horaria.", style='Caption')

    p_ans2 = doc.add_paragraph()
    p_ans2.add_run("Análisis de Demanda por Trámite: ").bold = True
    p_ans2.add_run(analisis_2_2 if analisis_2_2 else "No se registraron observaciones.")

    if buf_fig_capacidad:
        doc.add_paragraph("\nLa Figura 6 compara la tasa total de llegada de los clientes frente a la capacidad máxima instalada del sistema de atención en ventanilla.")
        doc.add_picture(buf_fig_capacidad, width=Inches(5.0))
        doc.add_paragraph("Figura 6. Curva de Arribos frente al Umbral de Capacidad Operativa.", style='Caption')
        
    p_ans3 = doc.add_paragraph()
    p_ans3.add_run("Análisis Diagnóstico de Capacidad: ").bold = True
    p_ans3.add_run(analisis_2_3 if analisis_2_3 else "No se registraron observaciones.")
    
    # SECCIÓN 2
    doc.add_page_break()
    doc.add_heading('2. Diseño y Parametrización del Gemelo Digital (FlexSim)', level=1)
    doc.add_paragraph(f"El proceso lógico fue estructurado mediante un diagrama BPMN (Enlace de referencia: {link_bpmn if link_bpmn else 'N/A'}). Las etapas específicas que componen el flujo en el simulador se describen detalladamente en la Tabla 3.")
    
    if imagen_bpmn:
        imagen_bpmn.seek(0)
        doc.add_picture(imagen_bpmn, width=Inches(5.5))
        doc.add_paragraph("Figura 7. Diagrama BPMN del proceso As-Is.", style='Caption')

    doc.add_heading('Tabla 3. Descripción de Etapas del Proceso', level=2)
    add_df_to_doc(df_bpmn[df_bpmn['Nombre_Etapa'] != ""], doc)
    
    doc.add_paragraph("\nLa Tabla 4 presenta la parametrización técnica ingresada en FlexSim, incluyendo reglas de ruteo, distribuciones estadísticas y asignación de recursos operativos.")
    doc.add_heading('Tabla 4. Parámetros Técnicos del Gemelo Digital', level=2)
    add_df_to_doc(df_params[df_params['Nombre_Etapa'] != ""], doc)
    
    if imagen_flexsim:
        doc.add_paragraph("\nLa Figura 8 muestra la construcción física (Layout) del modelo dentro del entorno virtual.")
        imagen_flexsim.seek(0)
        doc.add_picture(imagen_flexsim, width=Inches(5.5))
        doc.add_paragraph("Figura 8. Layout y entorno 3D del Gemelo Digital en FlexSim.", style='Caption')
    
    p_notas = doc.add_paragraph()
    p_notas.add_run("Observaciones sobre la construcción del modelo: ").bold = True
    p_notas.add_run(notas_flexsim if notas_flexsim else "No se registraron observaciones.")
    
    doc.add_paragraph("\nPara validar la eficiencia del sistema, se definieron metas operativas específicas estructuradas en la Tabla 5.")
    doc.add_heading('Tabla 5. Definición de KPIs y Umbrales Semafóricos', level=2)
    add_df_to_doc(df_kpis[df_kpis['Nombre_Indicador'] != ""], doc)
    
    # SECCIÓN 3
    doc.add_page_break()
    doc.add_heading('3. Ejecución de la Simulación y Triangulación de Resultados', level=1)
    
    doc.add_paragraph("Durante la corrida del modelo computacional, se extrajeron evidencias gráficas del comportamiento dinámico de las entidades y de los tableros de control.")
    if imagen_sim_corrida:
        imagen_sim_corrida.seek(0)
        doc.add_picture(imagen_sim_corrida, width=Inches(5))
        doc.add_paragraph("Figura 9. Ejecución del Gemelo Digital en tiempo real.", style='Caption')
    if imagen_dashboard_kpi:
        imagen_dashboard_kpi.seek(0)
        doc.add_picture(imagen_dashboard_kpi, width=Inches(5))
        doc.add_paragraph("Figura 10. Dashboard estadístico obtenido en FlexSim.", style='Caption')
        
    doc.add_paragraph("\nLa Tabla 6 consolida las desviaciones halladas entre la teoría y la simulación y activa los semáforos de advertencia.")
    doc.add_heading('Tabla 6. Triangulación y Evaluación de KPIs', level=2)
    add_df_to_doc(df_triangulacion[df_triangulacion['Indicador_Clave'] != ""], doc)
    
    doc.add_heading('4. Conclusiones y Propuestas de Intervención', level=1)
    
    p_diag = doc.add_paragraph()
    p_diag.add_run("Diagnóstico Cuantitativo y Cualitativo Final: ").bold = True
    p_diag.add_run(diagnostico_equipo if diagnostico_equipo else "No se registraron observaciones.")
    
    p_rec = doc.add_paragraph()
    p_rec.add_run("\nRecomendaciones de Mejora Continua: ").bold = True
    p_rec.add_run(recomendaciones_equipo if recomendaciones_equipo else "No se registraron propuestas.")

    # SECCIÓN 5: VALORACIÓN DEL APRENDIZAJE (NUEVO EN WORD)
    doc.add_page_break()
    doc.add_heading('5. Valoración del Aprendizaje y Autoevaluación de Competencias', level=1)
    doc.add_paragraph("A continuación, se consolidan las calificaciones obtenidas por el equipo consultor en la autoevaluación de competencias desarrolladas durante el proyecto:")
    
    df_eval_doc = pd.DataFrame({
        "Competencia Analizada": [
            "Modelamiento Matemático y Analítico",
            "Construcción de Gemelo Digital (FlexSim)",
            "Triangulación e Interpretación de Datos",
            "Pensamiento Sistémico y Propuestas To-Be"
        ],
        "Calificación (1 a 5)": [
            f"{eval_matematica} / 5",
            f"{eval_flexsim} / 5",
            f"{eval_triangulacion} / 5",
            f"{eval_sistemico} / 5"
        ]
    })
    add_df_to_doc(df_eval_doc, doc)

    doc.add_paragraph("\nReflexión Metacognitiva del Equipo:")
    
    p_ref1 = doc.add_paragraph()
    p_ref1.add_run("• Principales Aprendizajes: ").bold = True
    p_ref1.add_run(reflexion_aprensizaje if reflexion_aprensizaje else "No se registraron comentarios.")

    p_ref2 = doc.add_paragraph()
    p_ref2.add_run("• Retos y Dificultades Superadas: ").bold = True
    p_ref2.add_run(dificultades_superadas if dificultades_superadas else "No se registraron comentarios.")

    p_ref3 = doc.add_paragraph()
    p_ref3.add_run("• Transferencia y Aplicabilidad Futura: ").bold = True
    p_ref3.add_run(aplicaciones_futuras if aplicaciones_futuras else "No se registraron comentarios.")
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

st.download_button(
    label="📥 Generar y Descargar Documento Word Completo (.docx)",
    data=generar_word(),
    file_name="Informe_Consultoria_Bitacora_Completo.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
