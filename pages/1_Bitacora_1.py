import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import json

# Configuración inicial de la página
st.set_page_config(
    page_title="Caso 1: Va a volver a suceder otra vez", 
    page_icon="🎓", 
    layout="wide"
)

# ==============================================================================
# SISTEMA DE RESPALDO: GUARDAR / CARGAR BORRADOR (.JSON)
# ==============================================================================
st.sidebar.title("🛡️ Respaldo de Seguridad")
st.sidebar.markdown("Usa esta sección para evitar perder tu información en caso de cortes de internet.")

# Cargar Borrador
borrador_subido = st.sidebar.file_uploader("📂 Cargar Borrador (.json)", type=["json"], key="borrador_loader")

if borrador_subido is not None and "borrador_cargado" not in st.session_state:
    try:
        datos_borrador = json.load(borrador_subido)
        for key, value in datos_borrador.items():
            st.session_state[key] = value
        st.session_state["borrador_cargado"] = True
        st.sidebar.success("✅ ¡Borrador cargado con éxito!")
        st.rerun()
    except Exception as e:
        st.sidebar.error(f"Error al cargar borrador: {e}")

# Helper para obtener valores guardados en session_state o valor por defecto
def get_val(key, default=""):
    return st.session_state.get(key, default)

# ==============================================================================
# ENCABEZADO PRINCIPAL
# ==============================================================================
st.title("🏭 Modelación y Simulación en Ingeniería Industrial")
st.subheader("Plataforma de Bitácoras Web | Caso 1: ...va a volver a suceder otra vez")
st.caption("Docente: Ing. Mag. Néstor Castañeda")
st.markdown("---")

# ==============================================================================
# BLOQUE 1: HEADER DE CONTROL E INTEGRANTES
# ==============================================================================
st.markdown("### 📌 Bloque 1: Control de Equipo y Roles")

col1, col2 = st.columns(2)

with col1:
    estudiante_1 = st.text_input("📊 Estudiante 1 (Analista de Datos/Simulación):", value=get_val("k_est1"), key="k_est1", placeholder="Nombre completo")
    correo_1 = st.text_input("✉️ Correo Estudiante 1:", value=get_val("k_cor1"), key="k_cor1", placeholder="ejemplo@correo.edu.co")
    st.markdown("<br>", unsafe_allow_html=True)
    estudiante_2 = st.text_input("⚙️ Estudiante 2 (Analista de Procesos/Modelación):", value=get_val("k_est2"), key="k_est2", placeholder="Nombre completo")
    correo_2 = st.text_input("✉️ Correo Estudiante 2:", value=get_val("k_cor2"), key="k_cor2", placeholder="ejemplo@correo.edu.co")

with col2:
    grupo = st.text_input("👥 Grupo / Paralelo:", value=get_val("k_grupo"), key="k_grupo", placeholder="Ej: Grupo 01")
    fecha_sesion = st.text_input("📅 Fecha / Sesión:", value=get_val("k_fecha", "Caso 1 - La Cafetería"), key="k_fecha")

st.markdown("---")

# ==============================================================================
# BLOQUE 2: DEFINICIÓN DEL PROBLEMA Y DIAGRAMA DE ÁRBOL
# ==============================================================================
st.markdown("### 🎯 Bloque 2: Definición del Problema y Diagnóstico")

st.markdown("#### 📝 2.1 Definición General del Problema")
def_problema = st.text_area(
    "Describe la problemática central operativamente identificada en la cafetería:",
    value=get_val("k_prob"), key="k_prob",
    placeholder="Ej: Congestión crítica en hora pico con tiempos de espera excesivos...",
    height=90
)

st.markdown("#### 🌳 2.2 Árbol de Causas y Efectos")
col_img1, col_img2 = st.columns(2)
with col_img1:
    link_arbol = st.text_input("🔗 Enlace al Diagrama de Árbol:", value=get_val("k_link_arb"), key="k_link_arb")
with col_img2:
    img_arbol = st.file_uploader("🖼️ Adjuntar Imagen del Árbol de Problemas:", type=["png", "jpg", "jpeg"], key="arbol_img")

col_c1, col_c2 = st.columns(2)
with col_c1:
    causas_raiz = st.text_area("🔍 Causas Raíz (Estructurales/Operativas):", value=get_val("k_causas"), key="k_causas", placeholder="Ej: 1. Suposición de capacidad determinista por el asesor...", height=120)
with col_c2:
    efectos_sistema = st.text_area("⚠️ Efectos en el Sistema:", value=get_val("k_efectos"), key="k_efectos", placeholder="Ej: 1. Fila que crece de forma ilimitada. 2. Clientes molestos...", height=120)

st.markdown("---")

# ==============================================================================
# BLOQUE 3: MODELO AS-IS (DIAGRAMA BPMN)
# ==============================================================================
st.markdown("### ⚙️ Bloque 3: MODELO AS-IS (Diagrama de Flujo del Proceso Actual)")

col_b_link, col_b_img = st.columns(2)
with col_b_link:
    link_bpmn = st.text_input("🔗 Enlace al Diagrama BPMN AS-IS:", value=get_val("k_link_bpmn"), key="k_link_bpmn")
with col_b_img:
    img_bpmn = st.file_uploader("🖼️ Adjuntar Imagen del Diagrama BPMN AS-IS:", type=["png", "jpg", "jpeg"], key="bpmn_img")

descripcion_paso_a_paso = st.text_area(
    "📜 Descripción detallada de las actividades de proceso (AS-IS):",
    value=get_val("k_desc_bpmn"), key="k_desc_bpmn",
    placeholder="Paso 1: Arribo de entidades. Paso 2: Espera en fila única. Paso 3: Atención por cajero...",
    height=120
)

st.markdown("---")

# ==============================================================================
# BLOQUE 4: FUNDAMENTACIÓN TEÓRICA Y MARCO DE COLAS
# ==============================================================================
st.markdown("### 📚 Bloque 4: Fundamentación Teórica y Marco del Modelo")

col_th1, col_th2 = st.columns(2)
with col_th1:
    opciones_mod = ["Teoría de Colas - Modelo M/M/c", "Teoría de Colas - Modelo M/M/1", "Modelo M/M/c/K"]
    idx_mod = opciones_mod.index(get_val("k_mod_sel")) if get_val("k_mod_sel") in opciones_mod else 0
    modelo_seleccionado = st.selectbox("Estructura Teórica:", opciones_mod, index=idx_mod, key="k_mod_sel")
with col_th2:
    notacion_kendall = st.text_input("Notación de Kendall:", value=get_val("k_kendall", "M/M/c/FIFO/∞/∞"), key="k_kendall")

col_d1, col_d2, col_d3 = st.columns(3)
with col_d1:
    input_lambda = st.number_input("Tasa de llegada (λ) [cl/h]:", min_value=0.1, value=float(get_val("k_lambda", 30.0)), key="k_lambda")
with col_d2:
    input_mu = st.number_input("Tasa de servicio (μ) [cl/h]:", min_value=0.1, value=float(get_val("k_mu", 15.0)), key="k_mu")
with col_d3:
    input_c_actual = st.number_input("Número de servidores (c):", min_value=1, value=int(get_val("k_c", 2)), key="k_c")

referencias_bibliograficas = st.text_area(
    "📖 Fuentes y Referencias Bibliográficas Técnicas (APA / IEEE):",
    value=get_val("k_refs"), key="k_refs",
    placeholder="Ej:\n1. Taha, H. A. (2017). Investigación de operaciones (10a. ed.). Pearson.",
    height=100
)

st.markdown("---")

# ==============================================================================
# BLOQUE 5: CÓDIGO EN COLAB Y ANÁLISIS DEL ASESOR
# ==============================================================================
st.markdown("### 🧪 Bloque 5: Experimentación en Colab, Comparación y Análisis")

col_colab1, col_colab2 = st.columns(2)
with col_colab1:
    link_colab = st.text_input("🔗 Enlace al Notebook de Google Colab (.ipynb):", value=get_val("k_link_colab"), key="k_link_colab")
with col_colab2:
    img_colab = st.file_uploader("🖼️ Pantallazo del Código y Resultados en Colab:", type=["png", "jpg", "jpeg"], key="colab_img")

col_comp1, col_comp2 = st.columns(2)
with col_comp1:
    asesor_fixed = st.text_area("📌 Diagnóstico del Ingeniero Asesor (Datos Fijos):", value=get_val("k_asesor"), key="k_asesor", placeholder="Ej: El Asesor calculó c=2 asumiendo rho=100% y Wq=0...", height=130)
with col_comp2:
    modelo_colas_res = st.text_area("📊 Resultados del Modelo de Colas en Colab (M/M/c):", value=get_val("k_colab_res"), key="k_colab_res", placeholder="Ej: Con c=2 y variabilidad, la fila Wq y Lq tienden al infinito...", height=130)

st.markdown("---")

# ==============================================================================
# BLOQUE 6: SIMULACIÓN EN FLEXSIM Y TRIANGULACIÓN (NUEVA SECCIÓN DE INPUTS Y KPIS)
# ==============================================================================
st.markdown("### 🎮 Bloque 6: Simulación de Eventos Discretos en FlexSim 3D")

st.markdown("#### ⚙️ 6.1 Configuración de Parámetros de Entrada (Inputs en FlexSim)")
col_fx_in1, col_fx_in2 = st.columns(2)
with col_fx_in1:
    fx_inputs_llegadas = st.text_area(
        "📥 Distribución y Tasa de Llegadas (Inter-arrival time):", 
        value=get_val("k_fx_in_llegadas"), key="k_fx_in_llegadas",
        placeholder="Ej: Distribución Exponencial con media de 2 minutos por cliente (λ = 30 cl/h)...",
        height=100
    )
with col_fx_in2:
    fx_inputs_servicio = st.text_area(
        "🛠️ Distribución y Tiempo de Servicio por Cajero (Process time):", 
        value=get_val("k_fx_in_servicio"), key="k_fx_in_servicio",
        placeholder="Ej: Distribución Exponencial con media de 4 minutos por cliente (μ = 15 cl/h). Número de cajeros = 2...",
        height=100
    )

st.markdown("#### 🎯 6.2 Variables de Respuesta y Métricas Clave (KPIs u Outputs)")
fx_kpis_definidos = st.text_area(
    "📊 Define las Variables de Respuesta (KPIs) a monitorear en el Dashboard de FlexSim:",
    value=get_val("k_fx_kpis"), key="k_fx_kpis",
    placeholder="Ej: 1. Tiempo promedio en fila (Wq). 2. Longitud máxima de la fila (Lq). 3. % de Utilización de los cajeros (rho). 4. Total de salidas...",
    height=100
)

st.markdown("#### 🖼️ 6.3 Evidencias y Resultados Visuales de FlexSim")
col_fx_img1, col_fx_img2 = st.columns(2)
with col_fx_img1:
    link_flexsim_fsm = st.text_input("🔗 Enlace al archivo (.fsm):", value=get_val("k_link_fsm"), key="k_link_fsm")
    img_flexsim_model = st.file_uploader("🖼️ Pantallazo del Modelo 3D (FlexSim):", type=["png", "jpg", "jpeg"], key="fx_model_img")
with col_fx_img2:
    img_flexsim_dashboard = st.file_uploader("📊 Pantallazo del Dashboard de KPIs (FlexSim):", type=["png", "jpg", "jpeg"], key="fx_dash_img")

st.markdown("#### 🔬 6.4 Triangulación de Resultados y Análisis Crítico")
triangulacion_analitica = st.text_area("🔬 Triangulación de Resultados (Asesor vs Colab vs FlexSim):", value=get_val("k_triang"), key="k_triang", placeholder="Ej: Asesor (Wq=0), Colab (Wq->Infinito), FlexSim (Queue desbordada)...", height=150)
analisis_ingenieria = st.text_area("🔬 Análisis Crítico y Dictamen de Ingeniería Industrial:", value=get_val("k_analisis"), key="k_analisis", placeholder="Ej: Se demuestra la falacia del asesor al ignorar la variabilidad...", height=120)
propuesta_tobe = st.text_area("💡 6.5 Recomendaciones y Escenario TO-BE:", value=get_val("k_tobe"), key="k_tobe", placeholder="Ej: Se requiere c=3 para reducir rho al 66% y garantizar estabilidad...", height=120)

st.markdown("---")

# ==============================================================================
# BOTÓN EN LA BARRA LATERAL PARA DESCARGAR EL BORRADOR JSON
# ==============================================================================
estado_actual = {
    "k_est1": estudiante_1, "k_cor1": correo_1, "k_est2": estudiante_2, "k_cor2": correo_2,
    "k_grupo": grupo, "k_fecha": fecha_sesion, "k_prob": def_problema, "k_link_arb": link_arbol,
    "k_causas": causas_raiz, "k_efectos": efectos_sistema, "k_link_bpmn": link_bpmn,
    "k_desc_bpmn": descripcion_paso_a_paso, "k_mod_sel": modelo_seleccionado,
    "k_kendall": notacion_kendall, "k_lambda": input_lambda, "k_mu": input_mu, "k_c": input_c_actual,
    "k_refs": referencias_bibliograficas, "k_link_colab": link_colab, "k_asesor": asesor_fixed,
    "k_colab_res": modelo_colas_res, "k_fx_in_llegadas": fx_inputs_llegadas, "k_fx_in_servicio": fx_inputs_servicio,
    "k_fx_kpis": fx_kpis_definidos, "k_link_fsm": link_flexsim_fsm, "k_triang": triangulacion_analitica,
    "k_analisis": analisis_ingenieria, "k_tobe": propuesta_tobe
}

json_borrador = json.dumps(estado_actual, indent=4)
st.sidebar.download_button(
    label="💾 Descargar Borrador Backup (.json)",
    data=json_borrador,
    file_name=f"Borrador_Caso1_{grupo if grupo else 'Sistemas'}.json",
    mime="application/json"
)

# ==============================================================================
# EXPORTADOR OFICIAL A DOCUMENTO WORD (.DOCX) - INCLUYE INPUTS Y KPIS DE FLEXSIM
# ==============================================================================
def generar_word_oficial():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    NAVY = RGBColor(0, 51, 102)
    STEEL = RGBColor(70, 130, 180)
    WHITE = RGBColor(255, 255, 255)
    HEX_NAVY = "003366"
    fig_counter = 1

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

    def set_table_borders(table):
        tblPr = table._element.xpath('w:tblPr')
        if tblPr:
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders)

    def add_sec_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY

    def add_link_to_doc(url, label="🔗 Enlace asociado: "):
        if url:
            p = doc.add_paragraph()
            p.add_run(label).font.bold = True
            r = p.add_run(url)
            r.font.color.rgb = STEEL
            r.font.underline = True

    def add_image_with_caption(img_obj, caption_text):
        nonlocal fig_counter
        if img_obj is not None:
            try:
                img_bytes = io.BytesIO(img_obj.getvalue())
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_bytes, width=Inches(5.5))
                
                p_cap = doc.add_paragraph(f"Figura {fig_counter}: {caption_text}")
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.runs[0].font.size = Pt(8.5)
                p_cap.runs[0].font.italic = True
                p_cap.runs[0].font.color.rgb = STEEL
                fig_counter += 1
            except Exception as e:
                doc.add_paragraph(f"[Error al cargar la imagen: {e}]")

    # ENCABEZADO PERSONALIZADO
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_main = title_p.add_run("Caso 1: ...va a volver a suceder otra vez\n")
    r_main.font.size = Pt(16)
    r_main.font.bold = True
    r_main.font.color.rgb = NAVY
    
    r_sub = title_p.add_run("Modelación y Simulación\n")
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True
    
    r_doc = title_p.add_run("Ing. Mag. Néstor Castañeda")
    r_doc.font.size = Pt(11)
    r_doc.font.italic = True

    # BLOQUE 1: Tabla de Roles
    add_sec_header("Bloque 1: Control de Equipo e Identificación de Roles")
    tbl1 = doc.add_table(rows=3, cols=4)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl1)
    
    for i, h in enumerate(["Rol de Ingeniería", "Estudiante", "Correo", "Grupo / Fecha"]):
        c = tbl1.cell(0, i)
        set_cell_background(c, HEX_NAVY)
        p = c.paragraphs[0].add_run(h)
        p.font.bold = True
        p.font.color.rgb = WHITE

    tbl1.cell(1, 0).paragraphs[0].add_run("Analista de Datos / Colab")
    tbl1.cell(1, 1).paragraphs[0].add_run(estudiante_1 if estudiante_1 else "N/A")
    tbl1.cell(1, 2).paragraphs[0].add_run(correo_1 if correo_1 else "N/A")
    tbl1.cell(1, 3).paragraphs[0].add_run(f"{grupo} - {fecha_sesion}")

    tbl1.cell(2, 0).paragraphs[0].add_run("Analista FlexSim")
    tbl1.cell(2, 1).paragraphs[0].add_run(estudiante_2 if estudiante_2 else "N/A")
    tbl1.cell(2, 2).paragraphs[0].add_run(correo_2 if correo_2 else "N/A")
    tbl1.cell(2, 3).paragraphs[0].add_run(f"{grupo} - {fecha_sesion}")

    doc.add_paragraph("Tabla 1: Asignación de Roles.").runs[0].font.italic = True

    # BLOQUE 2: Árbol de Problemas
    add_sec_header("Bloque 2: Definición del Problema y Árbol de Causas")
    doc.add_paragraph(f"Problema Central: {def_problema}")
    add_link_to_doc(link_arbol, "🔗 Enlace al Diagrama de Árbol: ")
    add_image_with_caption(img_arbol, "Diagrama de Árbol de Causas y Efectos.")
    
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl2)
    tbl2.cell(0, 0).paragraphs[0].add_run("Causas Raíz").font.bold = True
    tbl2.cell(0, 1).paragraphs[0].add_run("Efectos").font.bold = True
    for c in range(2): set_cell_background(tbl2.cell(0, c), HEX_NAVY)
    tbl2.cell(1, 0).paragraphs[0].add_run(causas_raiz)
    tbl2.cell(1, 1).paragraphs[0].add_run(efectos_sistema)
    doc.add_paragraph("Tabla 2: Matriz Analítica de Causas y Efectos.").runs[0].font.italic = True

    # BLOQUE 3 & 4: BPMN y Teoría de Colas
    add_sec_header("Bloque 3 & 4: Procesos AS-IS y Parámetros del Modelo")
    add_link_to_doc(link_bpmn, "🔗 Enlace al Diagrama BPMN: ")
    add_image_with_caption(img_bpmn, "Mapeo del proceso AS-IS en notación BPMN.")
    doc.add_paragraph(descripcion_paso_a_paso)
    doc.add_paragraph(f"Modelo Teórico a evaluar: {modelo_seleccionado} | Notación de Kendall: {notacion_kendall}").runs[0].font.bold = True

    tbl3 = doc.add_table(rows=4, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl3)
    for i, h in enumerate(["Parámetro", "Valor", "Unidad"]):
        c = tbl3.cell(0, i)
        set_cell_background(c, HEX_NAVY)
        c.paragraphs[0].add_run(h).font.color.rgb = WHITE
    params_in = [("Tasa de llegada (λ)", str(input_lambda), "cl/h"), ("Tasa de servicio (μ)", str(input_mu), "cl/h"), ("Servidores (c)", str(input_c_actual), "cajeras")]
    for idx, (p_n, p_v, p_u) in enumerate(params_in, start=1):
        tbl3.cell(idx, 0).paragraphs[0].add_run(p_n)
        tbl3.cell(idx, 1).paragraphs[0].add_run(p_v)
        tbl3.cell(idx, 2).paragraphs[0].add_run(p_u)
    doc.add_paragraph("Tabla 3: Parámetros del Modelo Teoría de Colas.").runs[0].font.italic = True

    # BLOQUE 5: Colab
    add_sec_header("Bloque 5: Experimentación Analítica (Google Colab)")
    add_link_to_doc(link_colab, "🔗 Enlace al Notebook de Colab: ")
    add_image_with_caption(img_colab, "Código fuente y resultados de la evaluación analítica computacional.")

    # BLOQUE 6: FlexSim, Inputs, KPIs y Resultados
    add_sec_header("Bloque 6: Simulación de Eventos Discretos en FlexSim")
    
    doc.add_paragraph("Configuración de Parámetros de Entrada (Inputs en FlexSim):").runs[0].font.bold = True
    tbl_fx_in = doc.add_table(rows=2, cols=2)
    tbl_fx_in.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl_fx_in)
    tbl_fx_in.cell(0, 0).paragraphs[0].add_run("Tasa / Distribución de Llegadas").font.bold = True
    tbl_fx_in.cell(0, 1).paragraphs[0].add_run("Tasa / Distribución de Servicio").font.bold = True
    for c in range(2): set_cell_background(tbl_fx_in.cell(0, c), HEX_NAVY)
    tbl_fx_in.cell(1, 0).paragraphs[0].add_run(fx_inputs_llegadas)
    tbl_fx_in.cell(1, 1).paragraphs[0].add_run(fx_inputs_servicio)
    
    doc.add_paragraph("\nVariables de Respuesta y KPIs Monitoreados en FlexSim:").runs[0].font.bold = True
    doc.add_paragraph(fx_kpis_definidos)

    add_link_to_doc(link_flexsim_fsm, "🔗 Enlace al modelo .fsm (FlexSim): ")
    add_image_with_caption(img_flexsim_model, "Entorno de simulación de eventos discretos (FlexSim 3D).")
    add_image_with_caption(img_flexsim_dashboard, "Dashboard estadístico y recolección de KPIs.")

    tbl4 = doc.add_table(rows=2, cols=3)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl4)
    for i, h in enumerate(["Ing. Asesor (Fijo)", "Modelo Colab", "Simulación FlexSim"]):
        c = tbl4.cell(0, i)
        set_cell_background(c, HEX_NAVY)
        c.paragraphs[0].add_run(h).font.color.rgb = WHITE
    tbl4.cell(1, 0).paragraphs[0].add_run(asesor_fixed)
    tbl4.cell(1, 1).paragraphs[0].add_run(modelo_colas_res)
    tbl4.cell(1, 2).paragraphs[0].add_run(triangulacion_analitica)
    doc.add_paragraph("Tabla 4: Triangulación de Resultados Metodológicos.").runs[0].font.italic = True

    add_sec_header("Dictamen Crítico y Propuesta TO-BE")
    doc.add_paragraph(analisis_ingenieria)
    doc.add_paragraph(propuesta_tobe)

    add_sec_header("Referencias Bibliográficas")
    doc.add_paragraph(referencias_bibliograficas)

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# BOTÓN DE GENERACIÓN DE REPORTES
if st.button("📄 Generar y Descargar Bitácora Técnica Formateada (.docx)"):
    docx_file = generar_word_oficial()
    st.download_button(
        label="📥 Descargar Documento Word Oficial (.docx)",
        data=docx_file,
        file_name=f"Bitacora_Caso1_{grupo if grupo else 'Entregable'}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
